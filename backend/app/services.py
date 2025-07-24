from google import genai
from google.genai import types
from typing import List, Dict, Any, Optional, AsyncIterable, Union, BinaryIO
from google.generativeai.types import file_types
import logging
import httpx  # Add httpx for making API requests
import asyncio
import os
import mimetypes
from fastapi import HTTPException, status, UploadFile
from pathlib import Path
from datetime import datetime, timedelta
from sqlalchemy.orm import Session
import json
import uuid

from . import config, schemas, crud
from .crud import file_crud
from .models import FileMetadata

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# System instruction for the math chatbot persona
MATH_CHATBOT_SYSTEM_INSTRUCTION = """
You are an AI Math Chatbot designed to help students and professionals with mathematics problems.
Your capabilities include:

1. Solving math problems step-by-step, from basic arithmetic to advanced calculus, linear algebra, statistics, and more.
2. Explaining mathematical concepts clearly with examples.
3. Providing visual representations of mathematical concepts using LaTeX notation.
4. Helping debug mathematical code (Python, R, MATLAB, etc.).
5. Answering questions about mathematical history and applications.

Guidelines:
- Always show your work step-by-step when solving problems.
- Format mathematical expressions using LaTeX:
  * Use $...$ for inline math (e.g., $x^2 + 5$)
  * Use $$...$$ for display/block math (e.g., $$\\int_0^\\infty e^{-x} dx = 1$$)
  * Ensure all LaTeX expressions are properly escaped (e.g., \\int, \\sum, \\frac)
  * **When presenting a mathematical formula, equation or text with inline math (using $$...$$ for display/block math (e.g., $$\\int_0^\\infty e^{-x} dx = 1$$)) after a colon (e.g., "The formula is:"), always place the rendered math on a new line, using display math ($$...$$), and ensure it is centered with an appropriate amount of space before and after the math block. The structure after the colon must always be: new line: "$$", new line: "block math", new line: "$$".** For example:
    The Pythagorean theorem states:
    $$
    a^2 + b^2 = c^2
    $$
    This ensures the math is visually distinct and easy to read.
- **In the case of a text in a bullet point, When presenting a mathematical formula, equation or text with inline math (using $$...$$ for display/block math (e.g., $$\\int_0^\\infty e^{-x} dx = 1$$)) after a colon (e.g., "- The formula is:"), always place the rendered math on a new line, using display math ($$...$$), and ensure it is centered with an appropriate amount of space before and after the math block. The structure after the colon must always be: new line: "$$", new line: "block math", new line: "$$".**
- When presenting a text after a colon, always place the text on a new line. The structure after the colon must always be: new line: "text". For example:
    Here's what the theorem states:
    In a right-angled triangle, the square... 
- **Use the mathematical symbol '=>' wherever it is appropriate in your explanations, such as to indicate logical implication, result, or stepwise deduction in math reasoning.**
- **When you propose a final solution, always the final solution should be presented in a new line and in a box as follows:
    Final Answer:
    $$
    boxed{<solution>}
    $$
    where <solution> is the final solution to the problem.**
- Structure your response using clear markdown formatting:
  * Use headings (## for sections, ### for subsections) for organization
  * Use bullet points or numbered lists for steps
  * Use **bold** for emphasis and important concepts
  * Use code blocks (```python, ```r) for code examples
- When appropriate, explain the intuition behind mathematical concepts.
- If a question is ambiguous, ask for clarification.
- If you're unsure about an answer, acknowledge your uncertainty.
- Be encouraging and supportive, especially with students who are struggling.
- For complex problems, break down the solution into manageable parts.
- Maintain context from previous messages in the conversation to build on earlier explanations.
- Always at the end of your response, be engaging and ask a question or questions to the user regarding the topic of the conversation.

Remember that your goal is to help users understand mathematics, not just provide answers.
"""

# Configure the Gemini client using the new Google Gen AI SDK
try:
    if config.GEMINI_API_KEY:
        # Initialize the client with the API key
        client = genai.Client(api_key=config.GEMINI_API_KEY)

        # Set the model name from configuration
        model_name = config.get_settings().gemini_model_name

        # Log the model being used
        logger.info(f"Using Gemini model: {model_name}")

        # We use the client for both direct model calls and Files API
        model = client
        
        # Test that the model is accessible
        try:
            # Use the correct method to list models
            available_models = client.models.list()
            logger.info(f"Successfully connected to Gemini API. Models available.")
            
            # Don't check for exact model match - preview models may not be listed but still work
            # Just log that we're using the configured model
            logger.info(f"Using configured model: {model_name}")
                
        except Exception as e:
            logger.warning(f"Could not list available models: {e}. API key may be invalid or there might be connectivity issues.")
        
        logger.info("Gemini AI client configured successfully.")
    else:
        model = None
        logger.warning("GEMINI_API_KEY not found. AI functionality will be disabled.")
except Exception as e:
    model = None
    logger.error(f"Failed to configure Gemini client: {e}", exc_info=True)

def process_file_for_gemini(file_path: str, db: Session = None) -> Optional[Union[types.Part, file_types.File]]:
    """Process a file for inclusion in a Gemini API request.

    Returns either:
    - A Part object that can be included in the contents list for Gemini API (for small files)
    - A Gemini File object for large files processed through the Files API
    - None if processing fails

    Handles different file types appropriately:
    - PDF/Images: Passed directly to Gemini API (inline for <20MB, Files API for larger files)
    - TXT: Text is extracted and passed as text content
    - DOCX: Text is extracted using python-docx and passed as text content

    Args:
        file_path: Path to the file to process
        db: Optional database session for tracking Files API uploads
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return None

    try:
        # Get file size and mime type
        file_size = os.path.getsize(file_path)
        mime_type, _ = mimetypes.guess_type(file_path)

        if not mime_type:
            # Default to text/plain for unknown types
            mime_type = "text/plain"

        # Process different file types
        file_extension = os.path.splitext(file_path)[1].lower()

        # For text files, extract the text and return as text content
        if file_extension == '.txt' or mime_type == 'text/plain':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                text_content = f.read()
            logger.info(f"Extracted {len(text_content)} characters from text file: {file_path}")
            return types.Part(text=text_content)

        # For DOCX files, extract text using python-docx
        elif file_extension == '.docx' or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            text_content = extract_text_from_docx(file_path)
            logger.info(f"Extracted {len(text_content)} characters from DOCX file: {file_path}")
            return types.Part(text=text_content)

        # For PDF and image files, handle based on size
        else:
            # Check if file size exceeds the maximum allowed for inline data (20MB)
            max_inline_size = config.get_settings().max_file_size

            # For files under the inline size limit, use inline data approach
            if file_size <= max_inline_size:
                with open(file_path, "rb") as f:
                    file_data = f.read()

                logger.info(f"Processing {file_size} bytes of {mime_type} data inline from file: {file_path}")
                return types.Part(mime_type=mime_type, inline_data=file_data)

            # For larger files (up to 2GB), use the Gemini Files API
            else:
                # Check if model is configured (needed for Files API)
                if not model:
                    logger.error("Gemini model not initialized. Cannot use Files API.")
                    return None

                # Maximum size for Files API is 2GB
                max_files_api_size = 2 * 1024 * 1024 * 1024  # 2GB in bytes
                if file_size > max_files_api_size:
                    logger.error(f"File size ({file_size} bytes) exceeds maximum allowed for Files API ({max_files_api_size} bytes)")
                    return None

                logger.info(f"Uploading {file_size} bytes of {mime_type} data via Files API: {file_path}")

                # Extract file_id from the file path
                file_id = os.path.basename(file_path).split('.')[0]

                # Check if we already have a record for this file in the database
                if db:
                    existing_gemini_file = crud.get_gemini_file_by_file_id(db, file_id)
                    if existing_gemini_file and datetime.utcnow() < existing_gemini_file.expiry_time:
                        # If we have a valid record, use the existing Gemini file
                        logger.info(f"Using existing Gemini Files API record for file ID: {file_id}")
                        try:
                            # Check if we need to refresh the TTL
                            refreshed_file = refresh_gemini_file_if_needed(file_path, file_id, db)
                            if refreshed_file:
                                # Use the newly uploaded file if it was refreshed
                                logger.info(f"Using refreshed Gemini file: {refreshed_file.name}")
                                return refreshed_file
                                
                            # Get the file object from Gemini Files API using the stored name
                            file_obj = model.files.get(name=existing_gemini_file.gemini_file_name)
                            logger.info(f"Successfully retrieved existing file from Files API: {file_obj.name}")
                            return file_obj
                        except Exception as e:
                            logger.error(f"Error retrieving existing file from Files API: {e}", exc_info=True)
                            # If there's an error getting the file, the record might be stale
                            # We'll continue to upload a new file

                # Upload the file using the Files API
                try:
                    # Use the new SDK's files.upload method
                    uploaded_file = model.files.upload(
                        file=file_path,
                        config=dict(mime_type=mime_type)
                    )

                    logger.info(f"Successfully uploaded file to Gemini Files API. File ID: {uploaded_file.name}")

                    # Store the file ID and upload time in the database for TTL tracking
                    if db:
                        crud.create_gemini_file(
                            db=db,
                            file_id=file_id,
                            gemini_file_name=uploaded_file.name,
                            local_file_path=file_path,
                            mime_type=mime_type
                        )

                    return uploaded_file

                except Exception as e:
                    logger.error(f"Error uploading file to Gemini Files API: {e}", exc_info=True)
                    return None

    except Exception as e:
        logger.error(f"Error processing file {file_path}: {e}", exc_info=True)
        return None

def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from a DOCX file using python-docx.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content as a string
    """
    try:
        import docx
        doc = docx.Document(file_path)

        # Extract text from paragraphs
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        return '\n'.join(full_text)
    except ImportError:
        logger.error("python-docx library not installed. Cannot extract text from DOCX files.")
        return "[Error: Could not extract text from DOCX file due to missing dependencies]"
    except Exception as e:
        logger.error(f"Error extracting text from DOCX file {file_path}: {e}", exc_info=True)
        return f"[Error extracting text from DOCX file: {str(e)}]"

def generate_ai_response(chat_history: List[schemas.Message], 
                         user_message_content: str,
                         current_message_files: Optional[List[schemas.FileMetadataInfo]] = None, 
                         db: Session = None) -> str:
    """Generates an AI response using Gemini, considering chat history and files."""
    if not model: # `model` is the genai.Client instance
        logger.warning("Gemini client not configured. Returning a placeholder response.")
        return "AI functionality is currently unavailable."

    try:
        # Construct the full conversation history for Gemini
        gemini_history = []
        for msg in chat_history:
            message_parts = [types.Part(text=msg.content)]
            if msg.files: # msg.files is List[schemas.FileMetadataInfo]
                for file_meta in msg.files:
                    try:
                        if file_meta.content_type and (file_meta.content_type.startswith('text/') or file_meta.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'):
                            # Extract text for .txt and .docx
                            if not os.path.exists(file_meta.local_disk_path):
                                logger.warning(f"Local file {file_meta.local_disk_path} for historical message {msg.id} not found. Skipping file part.")
                                continue
                            
                            extracted_text = ""
                            if file_meta.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                                extracted_text = extract_text_from_docx(file_meta.local_disk_path)
                            else: # text/*
                                with open(file_meta.local_disk_path, 'r', encoding='utf-8', errors='replace') as f:
                                    extracted_text = f.read()
                            # Add a small preamble indicating the filename for text-based files
                            file_preamble = f"[Content from file: {file_meta.original_filename}]\\n"
                            message_parts.append(types.Part(text=file_preamble + extracted_text))
                        elif file_meta.gemini_api_file_id and file_meta.gemini_api_file_id.startswith("files/"): # Gemini File API URI
                            message_parts.append(types.Part(uri=file_meta.gemini_api_file_id, mime_type=file_meta.content_type))
                        else: # Inline file (PDF, image), file_id is the local_path
                            if not os.path.exists(file_meta.local_disk_path):
                                logger.warning(f"Local file {file_meta.local_disk_path} for historical message {msg.id} not found. Skipping file part.")
                                continue
                            with open(file_meta.local_disk_path, "rb") as f_data:
                                file_data = f_data.read()
                                message_parts.append(types.Part(mime_type=file_meta.content_type, inline_data=file_data))
                    except Exception as e:
                        logger.error(f"Error processing historical file {file_meta.original_filename} (ID: {file_meta.id}) for message {msg.id}: {e}", exc_info=True)
                        message_parts.append(types.Part(text=f"[Could not load content for file: {file_meta.original_filename}]"))
            
            gemini_history.append(types.Content(role=msg.role, parts=message_parts))

        # Now add the current user message and its files
        current_message_parts = [types.Part(text=user_message_content)]
        if current_message_files:
            for file_meta in current_message_files:
                try:
                    if file_meta.content_type and (file_meta.content_type.startswith('text/') or file_meta.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'):
                        if not os.path.exists(file_meta.local_disk_path):
                            logger.warning(f"Local file {file_meta.local_disk_path} for current message not found. Skipping file part.")
                            continue
                        extracted_text = ""
                        if file_meta.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                            extracted_text = extract_text_from_docx(file_meta.local_disk_path)
                        else: # text/*
                            with open(file_meta.local_disk_path, 'r', encoding='utf-8', errors='replace') as f:
                                extracted_text = f.read()
                        file_preamble = f"[Content from file: {file_meta.original_filename}]\\n"
                        current_message_parts.append(types.Part(text=file_preamble + extracted_text))
                    elif file_meta.gemini_api_file_id and file_meta.gemini_api_file_id.startswith("files/"): # Gemini File API URI
                        current_message_parts.append(types.Part(uri=file_meta.gemini_api_file_id, mime_type=file_meta.content_type))
                    else: # Inline file (PDF, image), file_id is the local_path
                        if not os.path.exists(file_meta.local_disk_path):
                            logger.warning(f"Local file {file_meta.local_disk_path} for current message not found. Skipping file part.")
                            continue
                        with open(file_meta.local_disk_path, "rb") as f_data:
                            file_data = f_data.read()
                            current_message_parts.append(types.Part(mime_type=file_meta.content_type, inline_data=file_data))
                except Exception as e:
                    logger.error(f"Error processing current file {file_meta.original_filename} (ID: {file_meta.id}): {e}", exc_info=True)
                    current_message_parts.append(types.Part(text=f"[Could not load content for file: {file_meta.original_filename}]"))
        
        # The full content for Gemini API call
        # The last element in gemini_history is the latest user message if chat_history included it.
        # If chat_history is from crud.get_messages_for_chat, it includes the user message that triggered this call.
        # We should ensure generate_ai_response is called AFTER user message (and its files) are saved.
        # So, gemini_history would be all messages up to BUT NOT INCLUDING the current one.
        # Then we append the current user message with its files.

        # Let's adjust: gemini_history should be history BEFORE current user turn.
        # The API call will be: model.generate_content(contents=[*gemini_history, types.Content(role='user', parts=current_message_parts)], ...)
        
        # Safety settings (optional, configure as needed)
        safety_settings = [
            {
                "category": "HARM_CATEGORY_HARASSMENT",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            },
            {
                "category": "HARM_CATEGORY_HATE_SPEECH",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            },
            {
                "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", 
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            },
            {
                "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                "threshold": "BLOCK_MEDIUM_AND_ABOVE"
            }
        ]
        
        generation_config = types.GenerationConfig(
            temperature=0.7,
            top_p=0.95,
            top_k=40,
            max_output_tokens=8192
        )

        # Make the API call using the client directly for non-streaming
        # Assuming `client` is the configured `genai.Client` instance
        gemini_model_to_use = client.generative_model(
            model_name=config.get_settings().gemini_model_name,
            system_instruction=MATH_CHATBOT_SYSTEM_INSTRUCTION,
            generation_config=types.GenerationConfig(
                temperature=0.7,
                top_p=0.95,
                top_k=40,
                max_output_tokens=8192
            ),
            safety_settings=[
                {
                    "category": "HARM_CATEGORY_HARASSMENT",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
                {
                    "category": "HARM_CATEGORY_HATE_SPEECH",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
                {
                    "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", 
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
                {
                    "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                }
            ]
        )

        response = gemini_model_to_use.generate_content(contents=[*gemini_history, types.Content(role='user', parts=current_message_parts)])

        # Handle potential errors or empty responses
        if not response.candidates or not response.candidates[0].content.parts:
            logger.warning("Gemini API returned no content or an unexpected response format.")
            return "I encountered an issue generating a response. Please try again."
        
        # Extract and return the text content from the response
        # Assuming the response is primarily text. For multimodal responses, this would need more handling.
        ai_response_text = ' '.join(part.text for part in response.candidates[0].content.parts if hasattr(part, 'text'))
        return ai_response_text.strip()

    except Exception as e:
        logger.error(f"Error generating AI response: {e}", exc_info=True)
        # More specific error mapping could be done here (e.g., for APIError subtypes)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Error generating AI response: {str(e)}")

async def save_uploaded_file(file: UploadFile) -> str:
    """Saves an uploaded file to the configured upload directory and returns its path.

    Args:
        file: The uploaded file object from FastAPI

    Returns:
        The path to the saved file

    Raises:
        HTTPException: If there's an error saving the file
    """
    upload_dir = Path(config.get_settings().upload_dir)
    upload_dir.mkdir(parents=True, exist_ok=True)  # Ensure directory exists

    # Create a unique filename (e.g., using UUID or timestamp + original name)
    # For now, using a simple approach. Ensure this is robust for production.
    # Potentially use file.filename directly if unique enough or prepend with chat_id/message_id later
    
    # Ensure filename is sanitized to prevent directory traversal or other issues.
    # Using a fixed prefix and only the base filename.
    base_filename = Path(file.filename).name 
    # A more unique name might be better, e.g. using uuid
    # unique_filename = f"{uuid.uuid4()}_{base_filename}\" 
    # For now, let's assume the router or another part of the system ensures some uniqueness context if needed
    # or that filenames might not be unique across all uploads but unique enough per message context later.
    # The current crud.get_file_by_id implies file_id (often part of filename) is the key.

    file_path = upload_dir / base_filename # Using base_filename as is.
                                         # Consider a more robust naming strategy for collisions.

    try:
        with open(file_path, "wb") as buffer:
            content = await file.read()  # Read file content
            buffer.write(content)        # Write to buffer
        logger.info(f"Successfully saved uploaded file: {file_path}")
    except Exception as e:
        logger.error(f"Error saving file {file.filename}: {e}", exc_info=True)
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail=f"Could not save file: {file.filename}")
    
    return str(file_path)

async def prepare_file_metadata_for_db(upload_file: UploadFile, db: Session) -> dict:
    """
    Saves the uploaded file, determines if it should use Gemini Files API,
    uploads to Gemini if necessary, and returns metadata for database storage.
    """
    local_path = await save_uploaded_file(upload_file)
    filename = Path(upload_file.filename).name # Sanitize/get base name
    content_type = upload_file.content_type
    
    # Get file size from the saved local file
    try:
        file_size = os.path.getsize(local_path)
    except OSError as e:
        logger.error(f"Could not get size of file {local_path}: {e}", exc_info=True)
        # Handle error, perhaps raise HTTPException or return a default/error state
        # For now, let's assume it will be small to avoid Files API upload if size is unknown
        file_size = 0 

    # Generate a unique file ID
    file_id = str(uuid.uuid4())
    
    # Determine if Gemini Files API should be used
    # Using max_file_size from config (e.g., 20MB for inline)
    max_inline_size = config.get_settings().max_file_size
    processing_method = "inline" if file_size <= max_inline_size else "files_api"
    
    gemini_api_file_id = None
    
    # Check if client is configured and Files API is needed
    if processing_method == "files_api" and model:
        # Max size for Files API is 2GB
        max_files_api_size = 2 * 1024 * 1024 * 1024  # 2GB in bytes
        if file_size > max_files_api_size:
            logger.error(f"File {filename} size ({file_size} bytes) exceeds maximum for Files API ({max_files_api_size} bytes). Will not upload.")
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File size exceeds {max_files_api_size/(1024*1024)}MB maximum allowed by Gemini Files API"
            )
        else:
            logger.info(f"File {filename} ({file_size} bytes) exceeds inline limit. Uploading to Gemini Files API.")
            try:
                # Upload to Gemini Files API
                gemini_sdk_file_obj = client.files.upload(
                    path=local_path, 
                    display_name=filename,
                    mime_type=content_type
                )
                gemini_api_file_id = gemini_sdk_file_obj.name  # This is the URI, e.g., "files/xxxxxxx"
                logger.info(f"Successfully uploaded {filename} to Gemini Files API. Gemini URI: {gemini_api_file_id}")
            
            except Exception as e:
                logger.error(f"Error uploading {filename} to Gemini Files API: {e}. Will fall back to inline processing.", exc_info=True)
                processing_method = "inline"
    
    # Create file metadata in database
    db_file_metadata = file_crud.create_file_metadata(
        db=db,
        file_id=file_id,
        original_filename=filename,
        content_type=content_type,
        size=file_size,
        local_disk_path=local_path,
        processing_method=processing_method
    )
    
    # If we uploaded to Gemini, update the metadata with Gemini info
    if gemini_api_file_id:
        db_file_metadata = file_crud.update_file_metadata_gemini_info(
            db=db,
            file_id=file_id,
            gemini_api_file_id=gemini_api_file_id
        )
    
    return {"id": db_file_metadata.id}

# Add a new function to refresh Files API TTL for files near expiry
def refresh_gemini_file_if_needed(file_path: str, file_id: str, db: Session = None) -> Optional[file_types.File]:
    """
    Check if a Gemini File API upload is near expiry and refresh it by re-uploading.
    
    Args:
        file_path: Path to the local file
        file_id: ID of the file
        db: Database session
        
    Returns:
        Updated Gemini File object if refreshed, or None if not needed/failed
    """
    # Only proceed if we have a database session and the model is configured
    if not db or not model:
        return None
        
    try:
        # Check if this file has a Gemini Files API record
        existing_gemini_file = crud.get_gemini_file_by_file_id(db, file_id)
        
        # If no record exists or file doesn't exist, return None
        if not existing_gemini_file or not os.path.exists(file_path):
            return None
            
        # Check if the file is approaching expiry (less than 10 hours remaining)
        # Gemini Files API has a 48-hour TTL
        expiry_threshold = datetime.utcnow() + timedelta(hours=10)
        
        if existing_gemini_file.expiry_time < expiry_threshold:
            logger.info(f"File {file_id} approaching expiry. Refreshing Files API TTL.")
            
            # Get mime type
            mime_type, _ = mimetypes.guess_type(file_path)
            if not mime_type:
                mime_type = "application/octet-stream"
                
            # Re-upload the file to reset the 48-hour TTL
            uploaded_file = model.files.upload(
                file=file_path,
                config=dict(mime_type=mime_type)
            )
            
            # Update the database record with the new file name and expiry time
            crud.update_gemini_file(
                db=db,
                file_id=file_id,
                gemini_file_name=uploaded_file.name
            )
            
            logger.info(f"Successfully refreshed Files API TTL for file {file_id}. New expiry in 48 hours.")
            return uploaded_file
            
        # File doesn't need refreshing yet
        return None
        
    except Exception as e:
        logger.error(f"Error refreshing Gemini Files API TTL for file {file_id}: {e}", exc_info=True)
        return None

async def generate_ai_response_stream(
    chat_id: str, 
    user_message_content: str,
    file_ids: Optional[List[str]],
    db: Session,
    queue: asyncio.Queue
) -> None:
    if not client:
        await queue.put(json.dumps({"error": "AI service not configured."}))
        await queue.put("[DONE]")
        return

    try:
        # 1. Save the user message and link files
        # The crud.create_chat_message now handles linking FileMetadata via file_ids
        # Note: chat_id from router is int, but crud might expect int if model.chat_id is int.
        # Assuming crud.create_chat_message handles string chat_id conversion if needed, or change here.
        try:
            user_msg_db = crud.create_chat_message(
                db=db, 
                chat_id=int(chat_id), # Ensure chat_id is int for DB op if schema expects int
                role="user", 
                content=user_message_content,
                file_ids=file_ids
            )
            logger.info(f"User message saved to DB with ID: {user_msg_db.id}")
            if file_ids:
                logger.info(f"Linked file_ids {file_ids} to message {user_msg_db.id}")
        except Exception as e:
            logger.error(f"Failed to save user message or link files: {e}", exc_info=True)
            await queue.put(json.dumps({"error": "Failed to save message."}))
            await queue.put("[DONE]")
            return

        # 2. Fetch complete chat history including the new message
        # The messages will have their .files relationship populated with FileMetadata objects
        chat_history_models = crud.get_messages_for_chat(db, chat_id=int(chat_id))

        # 3. Construct Gemini API contents list
        gemini_prompt_contents = [] 
        for msg_model in chat_history_models:
            # Fix: Create Part object correctly for message content
            content_part = types.Part(text=msg_model.content)
            message_parts = [content_part]
            if msg_model.files: # msg_model.files is List[models.FileMetadata]
                for fm_to_prepare in msg_model.files:
                    prepared_file_parts = await _prepare_single_file_for_gemini(fm_to_prepare, db)
                    if prepared_file_parts:
                        # Flatten the returned list of Parts (context + content)
                        message_parts.extend(prepared_file_parts)
            # Determine role for Gemini API (model vs user)
            gemini_role = "user" if msg_model.role == "user" else "model"
            gemini_prompt_contents.append(types.Content(role=gemini_role, parts=message_parts))

        # 4. AI response generation (streaming)
        logger.info(f"Sending request to Gemini with {len(gemini_prompt_contents)} contents for chat_id {chat_id}")
        # For debugging, can log parts of the prompt, carefully for PII if files are involved.
        # if gemini_prompt_contents and gemini_prompt_contents[-1].parts:
        #    logger.debug(f"Last prompt part content: {gemini_prompt_contents[-1].parts[0].text[:200]}...")
        
        response_stream = client.models.generate_content_stream(
            model=config.get_settings().gemini_model_name,
            contents=gemini_prompt_contents,
            config=types.GenerateContentConfig(
                system_instruction=MATH_CHATBOT_SYSTEM_INSTRUCTION,
                temperature=0.7,
                top_p=0.95,
                top_k=40,
                max_output_tokens=8192
            )
        )

        ai_response_content = ""
        for chunk in response_stream:
            if (
                chunk.candidates
                and getattr(chunk.candidates[0], "content", None) is not None
                and hasattr(chunk.candidates[0].content, "parts")
                and chunk.candidates[0].content.parts is not None
            ):
                for part in chunk.candidates[0].content.parts:
                    if hasattr(part, 'text'):
                        text_chunk = part.text
                        ai_response_content += text_chunk
                        await queue.put(json.dumps({"text": text_chunk}))
            else:
                logger.warning(
                    f"Skipping chunk with missing or empty content: finish_reason={getattr(chunk.candidates[0], 'finish_reason', None)}"
                )
                # Optionally, handle finish_reason or safety_ratings here if needed
        
        # Save AI's response to DB (no files attached to AI response in this flow)
        if ai_response_content:
            crud.create_chat_message(
                db=db, 
                chat_id=int(chat_id), 
                role="model", # Use 'model' to match DB constraint
                content=ai_response_content,
                file_ids=None # AI doesn't attach files in its response here
            )
            logger.info(f"AI response saved for chat_id {chat_id}")

    except Exception as e:
        logger.error(f"Error in AI response generation stream for chat {chat_id}: {e}", exc_info=True)
        error_detail = {"error": str(e), "text": "An error occurred while generating the AI response."}
        try:
            await queue.put(json.dumps(error_detail))
        except Exception as qe:
            logger.error(f"Failed to put error on queue: {qe}")
    finally:
        try:
            await queue.put("[DONE]")
        except Exception as qe:
            logger.error(f"Failed to put [DONE] on queue: {qe}")

# Remove old process_file_for_gemini, save_uploaded_file, prepare_file_metadata_for_db, etc.
# as their logic is now integrated or replaced by the new FileMetadata flow.
# The old generate_ai_response (non-streaming) might also be obsolete or need similar refactoring.

# refresh_gemini_file_if_needed needs to be adapted if it was using the old GeminiFile model.

def refresh_gemini_file_if_needed(fm: FileMetadata, db: Session) -> Optional[file_types.File]:
    """Refreshes a Gemini file if it's close to expiry or already expired, then returns the File object."""
    if not client or not fm.gemini_api_file_id or not fm.gemini_api_expiry_timestamp:
        return None # Cannot refresh

    # Check if nearing expiry (e.g., within last hour) or past expiry
    if datetime.utcnow() > (fm.gemini_api_expiry_timestamp - timedelta(hours=1)):
        logger.info(f"Gemini file {fm.gemini_api_file_id} for {fm.original_filename} is expiring soon or expired. Attempting re-upload.")
        try:
            # Re-upload from local_disk_path
            new_uploaded_gemini_file = client.files.upload(
                path=fm.local_disk_path,
                display_name=fm.original_filename,
                mime_type=fm.content_type
            )
            # Update DB
            updated_fm = file_crud.update_file_gemini_details(
                db=db, 
                file_id=fm.id, 
                gemini_api_file_id=new_uploaded_gemini_file.name, 
                gemini_api_upload_timestamp=datetime.utcnow()
            )
            logger.info(f"Refreshed Gemini file for {fm.original_filename}. New Gemini ID: {new_uploaded_gemini_file.name}")
            return new_uploaded_gemini_file
        except Exception as e:
            logger.error(f"Failed to re-upload/refresh Gemini file {fm.original_filename}: {e}", exc_info=True)
            # If re-upload fails, try to get the old one if it hasn't hard expired on Gemini side
            try:
                return client.files.get(name=fm.gemini_api_file_id)
            except:
                return None # Truly unavailable
    else:
        # Not expiring soon, just get the existing one
        try:
            return client.files.get(name=fm.gemini_api_file_id)
        except Exception as e:
            logger.warning(f"Could not get non-expiring Gemini file {fm.gemini_api_file_id}: {e}. It might have been deleted on Gemini side.")
            return None

# Refactored function to prepare a single FileMetadata object for Gemini
async def _prepare_single_file_for_gemini(fm: FileMetadata, db: Session) -> Optional[list]:
    if not client:
        logger.error("Gemini client not initialized.")
        return None

    # Explicitly check if the local file exists, as FileMetadata might be stale
    if not fm.local_disk_path or not os.path.exists(fm.local_disk_path):
        logger.error(f"Local file not found at path: {fm.local_disk_path} for FileMetadata ID: {fm.id}. Skipping this file.")
        return None

    try:
        context_part = types.Part(text=f"[File attached: {fm.original_filename}, type: {fm.content_type}]")
        # For text and docx, extract content and return as text Part
        if fm.content_type == 'text/plain':
            with open(fm.local_disk_path, 'r', encoding='utf-8', errors='replace') as f:
                text_content = f.read()
            return [context_part, types.Part(text=text_content)]
        elif fm.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            text_content = extract_text_from_docx(fm.local_disk_path)
            return [context_part, types.Part(text=text_content)]

        # For other types (PDF, images)
        if fm.processing_method == "inline":
            with open(fm.local_disk_path, "rb") as f_bytes:
                file_data = f_bytes.read()
            return [context_part, types.Part(inline_data={"data": file_data, "mime_type": fm.content_type})]
        elif fm.processing_method == "files_api":
            if fm.gemini_api_file_id and fm.gemini_api_expiry_timestamp and datetime.utcnow() < fm.gemini_api_expiry_timestamp:
                try:
                    gemini_file_obj = client.files.get(name=fm.gemini_api_file_id)
                    logger.info(f"Using existing Gemini File: {gemini_file_obj.name} for {fm.original_filename}")
                    return [context_part, gemini_file_obj]
                except Exception as e:
                    logger.warning(f"Failed to get existing Gemini file {fm.gemini_api_file_id} for {fm.original_filename}: {e}. Will re-upload.")
            logger.info(f"Uploading {fm.original_filename} ({fm.size} bytes) to Gemini Files API from {fm.local_disk_path}")
            uploaded_gemini_file = client.files.upload(
                path=fm.local_disk_path, 
                display_name=fm.original_filename,
                mime_type=fm.content_type
            )
            logger.info(f"Uploaded {fm.original_filename} to Gemini. Gemini File ID: {uploaded_gemini_file.name}")
            file_crud.update_file_gemini_details(
                db=db, 
                file_id=fm.id, 
                gemini_api_file_id=uploaded_gemini_file.name, 
                gemini_api_upload_timestamp=datetime.utcnow()
            )
            return [context_part, uploaded_gemini_file]
        else:
            logger.warning(f"Unknown processing_method '{fm.processing_method}' for file {fm.id}")
            return None

    except Exception as e:
        logger.error(f"Error preparing file {fm.id} ({fm.original_filename}) for Gemini: {e}", exc_info=True)
        return None